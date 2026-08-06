from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    NextPageTemplate,
    PageBreak,
    PageTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parent
CONTENT = ROOT / "report-content.txt"
PDF_OUT = ROOT / "MJC_AI_Campus_Agent_해커톤_보고서.pdf"
DOCX_OUT = ROOT / "MJC_AI_Campus_Agent_해커톤_보고서.docx"

NAVY = colors.HexColor("#071B33")
BLUE = colors.HexColor("#1677FF")
CYAN = colors.HexColor("#47C5FF")
INK = colors.HexColor("#172033")
MUTED = colors.HexColor("#5D6B82")
PALE = colors.HexColor("#EDF5FF")
LINE = colors.HexColor("#D8E4F2")
WHITE = colors.white


def register_fonts() -> None:
    candidates = [
        ("Malgun", Path("C:/Windows/Fonts/malgun.ttf")),
        ("MalgunBold", Path("C:/Windows/Fonts/malgunbd.ttf")),
    ]
    for name, path in candidates:
        if path.exists():
            pdfmetrics.registerFont(TTFont(name, str(path)))
    if "Malgun" not in pdfmetrics.getRegisteredFontNames():
        raise RuntimeError("맑은 고딕 글꼴을 찾을 수 없습니다.")


def parse_content() -> tuple[list[str], list[tuple[str, list[str]]]]:
    lines = [line.rstrip() for line in CONTENT.read_text(encoding="utf-8").splitlines()]
    header: list[str] = []
    sections: list[tuple[str, list[str]]] = []
    current_title: str | None = None
    current_lines: list[str] = []
    for line in lines:
        match = re.fullmatch(r"\[(.+)]", line.strip())
        if match:
            if current_title is not None:
                sections.append((current_title, current_lines))
            current_title = match.group(1)
            current_lines = []
        elif current_title is None:
            if line:
                header.append(line)
        else:
            current_lines.append(line)
    if current_title is not None:
        sections.append((current_title, current_lines))
    return header, sections


def footer(canvas, doc) -> None:
    if doc.page == 1:
        return
    canvas.saveState()
    canvas.setStrokeColor(LINE)
    canvas.setLineWidth(0.5)
    canvas.line(18 * mm, 14 * mm, 192 * mm, 14 * mm)
    canvas.setFont("Malgun", 7.8)
    canvas.setFillColor(MUTED)
    canvas.drawString(18 * mm, 9 * mm, "MJC AI Campus Agent · 코드한입조")
    canvas.drawRightString(192 * mm, 9 * mm, str(doc.page))
    canvas.restoreState()


def cover(canvas, doc) -> None:
    canvas.saveState()
    width, height = A4
    canvas.setFillColor(NAVY)
    canvas.rect(0, 0, width, height, fill=1, stroke=0)
    canvas.setFillColor(BLUE)
    canvas.rect(0, height - 13 * mm, width, 13 * mm, fill=1, stroke=0)
    canvas.setFillColor(CYAN)
    canvas.circle(width - 26 * mm, height - 34 * mm, 7 * mm, fill=1, stroke=0)
    canvas.setStrokeColor(colors.Color(1, 1, 1, alpha=0.12))
    canvas.setLineWidth(1)
    for x in range(-30, 240, 24):
        canvas.line(x * mm, 0, (x + 95) * mm, height)
    canvas.restoreState()


def build_pdf(header: list[str], sections: list[tuple[str, list[str]]]) -> None:
    register_fonts()
    doc = BaseDocTemplate(
        str(PDF_OUT),
        pagesize=A4,
        leftMargin=18 * mm,
        rightMargin=18 * mm,
        topMargin=19 * mm,
        bottomMargin=19 * mm,
        title="MJC AI Campus Agent 해커톤 보고서",
        author="코드한입조",
    )
    cover_frame = Frame(0, 0, A4[0], A4[1], id="cover", leftPadding=0, rightPadding=0, topPadding=0, bottomPadding=0)
    body_frame = Frame(18 * mm, 18 * mm, 174 * mm, 260 * mm, id="body", leftPadding=0, rightPadding=0, topPadding=0, bottomPadding=0)
    doc.addPageTemplates([
        PageTemplate(id="Cover", frames=[cover_frame], onPage=cover),
        PageTemplate(id="Body", frames=[body_frame], onPage=footer),
    ])

    styles = getSampleStyleSheet()
    title = ParagraphStyle("CoverTitle", parent=styles["Title"], fontName="MalgunBold", fontSize=28, leading=37, textColor=WHITE, alignment=TA_LEFT, spaceAfter=8 * mm)
    subtitle = ParagraphStyle("CoverSubtitle", parent=styles["Normal"], fontName="Malgun", fontSize=14, leading=22, textColor=colors.HexColor("#BCD8FF"), alignment=TA_LEFT)
    cover_meta = ParagraphStyle("CoverMeta", parent=styles["Normal"], fontName="Malgun", fontSize=10.5, leading=19, textColor=WHITE, alignment=TA_LEFT)
    h1 = ParagraphStyle("H1", parent=styles["Heading1"], fontName="MalgunBold", fontSize=19, leading=25, textColor=NAVY, spaceAfter=6 * mm, keepWithNext=True)
    kicker = ParagraphStyle("Kicker", parent=styles["Normal"], fontName="MalgunBold", fontSize=8.5, leading=12, textColor=BLUE, spaceAfter=2 * mm)
    body = ParagraphStyle("Body", parent=styles["BodyText"], fontName="Malgun", fontSize=9.5, leading=15.5, textColor=INK, alignment=TA_JUSTIFY, spaceAfter=4.1 * mm, wordWrap="CJK")
    small = ParagraphStyle("Small", parent=body, fontSize=8.4, leading=13, textColor=MUTED, alignment=TA_LEFT)
    metric = ParagraphStyle("Metric", parent=styles["Normal"], fontName="MalgunBold", fontSize=20, leading=24, textColor=BLUE, alignment=TA_CENTER)
    metric_label = ParagraphStyle("MetricLabel", parent=styles["Normal"], fontName="Malgun", fontSize=8.5, leading=12, textColor=MUTED, alignment=TA_CENTER)

    story = [Spacer(1, 58 * mm)]
    story.append(Paragraph("MJC AI<br/>Campus Agent", title))
    story.append(Paragraph("명지전문대학생을 위한<br/>신뢰 가능한 AI 캠퍼스 비서", subtitle))
    story.append(Spacer(1, 36 * mm))
    story.append(Paragraph("2026학년도 RISE사업단 AI 해커톤 경진대회", cover_meta))
    story.append(Spacer(1, 4 * mm))
    story.append(Paragraph("팀명  코드한입조", cover_meta))
    story.append(Paragraph("전자공학과 임채호  ·  AI게임소프트웨어학과 조영남", cover_meta))
    story.append(Spacer(1, 5 * mm))
    story.append(Paragraph("GitHub  github.com/hoo419/mjc-ai-hackathon-codeonebite-team", cover_meta))
    story.append(NextPageTemplate("Body"))
    story.append(PageBreak())

    for index, (section_title, paragraphs) in enumerate(sections):
        section_no = f"{index:02d}" if index else "00"
        story.append(Paragraph(f"SECTION {section_no}", kicker))
        story.append(Paragraph(section_title, h1))

        if section_title == "요약":
            metrics = [
                [Paragraph("246", metric), Paragraph("516", metric), Paragraph("14+1", metric), Paragraph("6", metric)],
                [Paragraph("고유 분반", metric_label), Paragraph("수업 세션", metric_label), Paragraph("서비스 API + Health", metric_label), Paragraph("사용자 화면", metric_label)],
            ]
            table = Table(metrics, colWidths=[43.5 * mm] * 4, rowHeights=[14 * mm, 10 * mm])
            table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, -1), PALE),
                ("BOX", (0, 0), (-1, -1), 0.7, LINE),
                ("INNERGRID", (0, 0), (-1, -1), 0.4, LINE),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
            ]))
            story.append(table)
            story.append(Spacer(1, 7 * mm))

        for p in paragraphs:
            if p.strip():
                story.append(Paragraph(p.replace("&", "&amp;"), body))

        if section_title == "5. 시스템 구조":
            flow = [["Next.js", "REST API", "FastAPI", "Service", "Repository", "JSON / PostgreSQL"]]
            t = Table(flow, colWidths=[30 * mm, 23 * mm, 26 * mm, 26 * mm, 31 * mm, 38 * mm], rowHeights=12 * mm)
            t.setStyle(TableStyle([
                ("FONTNAME", (0, 0), (-1, -1), "MalgunBold"),
                ("FONTSIZE", (0, 0), (-1, -1), 7.5),
                ("TEXTCOLOR", (0, 0), (-1, -1), NAVY),
                ("BACKGROUND", (0, 0), (-1, -1), PALE),
                ("BOX", (0, 0), (-1, -1), 0.6, BLUE),
                ("INNERGRID", (0, 0), (-1, -1), 0.4, LINE),
                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ]))
            story.append(Spacer(1, 2 * mm))
            story.append(t)
            story.append(Paragraph("데이터 저장 방식이 달라져도 Route와 Service는 동일한 계약을 유지한다.", small))

        if index < len(sections) - 1:
            story.append(PageBreak())

    doc.build(story)


def set_run_font(run, name: str, size: float, bold: bool = False, color: str = "172033") -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])


def build_docx(header: list[str], sections: list[tuple[str, list[str]]]) -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    normal = doc.styles["Normal"]
    normal.font.name = "맑은 고딕"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 18, "071B33", 16, 8),
        ("Heading 2", 13, "1677FF", 12, 6),
        ("Heading 3", 11, "071B33", 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = "맑은 고딕"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(90)
    r = p.add_run("MJC AI\nCampus Agent")
    set_run_font(r, "맑은 고딕", 30, True, "071B33")
    p = doc.add_paragraph()
    r = p.add_run("명지전문대학생을 위한 신뢰 가능한 AI 캠퍼스 비서")
    set_run_font(r, "맑은 고딕", 15, False, "1677FF")
    doc.add_paragraph("\n\n2026학년도 RISE사업단 AI 해커톤 경진대회")
    doc.add_paragraph("팀명  코드한입조")
    doc.add_paragraph("전자공학과 임채호 · AI게임소프트웨어학과 조영남")
    doc.add_paragraph("GitHub  https://github.com/hoo419/mjc-ai-hackathon-codeonebite-team")
    doc.add_page_break()

    for idx, (title, paragraphs) in enumerate(sections):
        doc.add_heading(title, level=1)
        for text in paragraphs:
            if text.strip():
                p = doc.add_paragraph(text)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if idx < len(sections) - 1:
            doc.add_page_break()

    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_p.add_run("MJC AI Campus Agent  ·  ")
    add_page_number(footer_p)
    for run in footer_p.runs:
        set_run_font(run, "맑은 고딕", 8, False, "5D6B82")
    doc.save(DOCX_OUT)


def main() -> None:
    header, sections = parse_content()
    if len(sections) != 11:
        raise ValueError(f"예상 섹션 수 11개, 실제 {len(sections)}개")
    build_pdf(header, sections)
    build_docx(header, sections)
    print(PDF_OUT)
    print(DOCX_OUT)


if __name__ == "__main__":
    main()
