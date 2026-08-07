from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    NextPageTemplate,
    PageBreak,
    PageTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parent
CONTENT = ROOT / "dev-report-content.txt"
PDF_OUT = ROOT / "MJC_AI_Campus_Agent_개발보고서.pdf"
DOCX_OUT = ROOT / "MJC_AI_Campus_Agent_개발보고서.docx"

NAVY = colors.HexColor("#071B33")
BLUE = colors.HexColor("#1677FF")
INK = colors.HexColor("#172033")
MUTED = colors.HexColor("#5D6B82")
PALE = colors.HexColor("#EDF5FF")
LINE = colors.HexColor("#D8E4F2")
WHITE = colors.white

Block = tuple[str, object]  # ("text", str) or ("table", list[list[str]])


def register_fonts() -> None:
    candidates = [
        ("Malgun", Path("C:/Windows/Fonts/malgun.ttf")),
        ("MalgunBold", Path("C:/Windows/Fonts/malgunbd.ttf")),
    ]
    for name, path in candidates:
        if path.exists():
            pdfmetrics.registerFont(TTFont(name, str(path)))
    if "Malgun" not in pdfmetrics.getRegisteredFontNames():
        raise RuntimeError("맑은 고딕 글꼴을 찾을 수 없습니다.")


BLOCK_MARKERS = {"TABLE", "BULLETS", "OLIST"}


def parse_content() -> tuple[list[str], list[tuple[str, list[Block]]]]:
    """report-content.txt와 같은 [섹션명] 마커 형식에 더해 다음을 처리한다:
    - [TABLE]..[/TABLE]: 파이프로 구분한 표
    - [BULLETS]..[/BULLETS]: 글머리 기호 목록 (한 줄에 항목 하나)
    - [OLIST]..[/OLIST]: 번호 매김 목록 (렌더링 시 1부터 새로 번호를 매김)
    - '##제목' 한 줄: 섹션 내 소제목(H2)
    """
    lines = [line.rstrip() for line in CONTENT.read_text(encoding="utf-8").splitlines()]
    header: list[str] = []
    sections: list[tuple[str, list[Block]]] = []
    current_title: str | None = None
    blocks: list[Block] = []
    para_lines: list[str] = []
    block_mode: str | None = None  # "TABLE" | "BULLETS" | "OLIST" | None
    block_rows: list[str | list[str]] = []

    def flush_paragraph() -> None:
        nonlocal para_lines
        if para_lines:
            blocks.append(("text", " ".join(para_lines).strip()))
            para_lines = []

    def flush_block() -> None:
        nonlocal block_mode, block_rows
        if block_mode == "TABLE":
            blocks.append(("table", block_rows))
        elif block_mode == "BULLETS":
            blocks.append(("bullets", block_rows))
        elif block_mode == "OLIST":
            blocks.append(("olist", block_rows))
        block_mode = None
        block_rows = []

    for line in lines:
        stripped = line.strip()

        if block_mode is not None:
            if stripped == f"[/{block_mode}]":
                flush_block()
            elif stripped:
                block_rows.append(stripped.split("|") if block_mode == "TABLE" else stripped)
            continue

        open_match = re.fullmatch(r"\[(TABLE|BULLETS|OLIST)]", stripped)
        if open_match:
            flush_paragraph()
            block_mode = open_match.group(1)
            block_rows = []
            continue

        section_match = re.fullmatch(r"\[(.+)]", stripped)
        h2_match = stripped.startswith("##")

        if section_match:
            flush_paragraph()
            if current_title is not None:
                sections.append((current_title, blocks))
            current_title = section_match.group(1)
            blocks = []
        elif current_title is None:
            if line:
                header.append(line)
        elif h2_match:
            flush_paragraph()
            blocks.append(("h2", stripped[2:].strip()))
        else:
            if not stripped:
                flush_paragraph()
            else:
                para_lines.append(stripped)

    flush_paragraph()
    if current_title is not None:
        sections.append((current_title, blocks))
    return header, sections


def footer(canvas, doc) -> None:
    if doc.page == 1:
        return
    canvas.saveState()
    canvas.setStrokeColor(LINE)
    canvas.setLineWidth(0.5)
    canvas.line(18 * mm, 14 * mm, 192 * mm, 14 * mm)
    canvas.setFont("Malgun", 7.8)
    canvas.setFillColor(MUTED)
    canvas.drawString(18 * mm, 9 * mm, "MJC AI Campus Agent · 개발 프로젝트 보고서")
    canvas.drawRightString(150 * mm, 9 * mm, "CodeOneBite Team | 2026")
    canvas.drawRightString(192 * mm, 9 * mm, str(doc.page))
    canvas.restoreState()


def cover(canvas, doc) -> None:
    canvas.saveState()
    width, height = A4
    canvas.setFillColor(WHITE)
    canvas.rect(0, 0, width, height, fill=1, stroke=0)
    canvas.setFillColor(MUTED)
    canvas.setFont("Malgun", 8.5)
    canvas.drawRightString(width - 18 * mm, height - 14 * mm, "MJC AI Campus Agent · 개발 프로젝트 보고서")
    canvas.drawCentredString(width / 2, 14 * mm, "CodeOneBite Team | 2026")
    canvas.restoreState()


def build_table(rows: list[list[str]]) -> Table:
    n_cols = len(rows[0])
    body_style = ParagraphStyle(
        "TableBody", fontName="Malgun", fontSize=8.3, leading=12.5, textColor=INK, wordWrap="CJK"
    )
    head_style = ParagraphStyle(
        "TableHead", fontName="MalgunBold", fontSize=8.3, leading=12.5, textColor=NAVY, wordWrap="CJK"
    )
    data = []
    for r_idx, row in enumerate(rows):
        style = head_style if r_idx == 0 else body_style
        data.append([Paragraph(cell.replace("&", "&amp;"), style) for cell in row])

    col_width = 174 * mm / n_cols
    table = Table(data, colWidths=[col_width] * n_cols, repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), PALE),
                ("BOX", (0, 0), (-1, -1), 0.6, LINE),
                ("INNERGRID", (0, 0), (-1, -1), 0.4, LINE),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]
        )
    )
    return table


def build_pdf(header: list[str], sections: list[tuple[str, list[Block]]]) -> None:
    register_fonts()
    doc = BaseDocTemplate(
        str(PDF_OUT),
        pagesize=A4,
        leftMargin=18 * mm,
        rightMargin=18 * mm,
        topMargin=19 * mm,
        bottomMargin=19 * mm,
        title="MJC AI Campus Agent 개발 프로젝트 보고서",
        author="CodeOneBite Team",
    )
    cover_frame = Frame(0, 0, A4[0], A4[1], id="cover", leftPadding=0, rightPadding=0, topPadding=0, bottomPadding=0)
    body_frame = Frame(18 * mm, 18 * mm, 174 * mm, 260 * mm, id="body", leftPadding=0, rightPadding=0, topPadding=0, bottomPadding=0)
    doc.addPageTemplates(
        [
            PageTemplate(id="Cover", frames=[cover_frame], onPage=cover),
            PageTemplate(id="Body", frames=[body_frame], onPage=footer),
        ]
    )

    styles = getSampleStyleSheet()
    kicker_cover = ParagraphStyle("KickerCover", parent=styles["Normal"], fontName="MalgunBold", fontSize=9.5, leading=14, textColor=BLUE, alignment=TA_CENTER)
    title = ParagraphStyle("CoverTitle", parent=styles["Title"], fontName="MalgunBold", fontSize=27, leading=34, textColor=NAVY, alignment=TA_CENTER, spaceAfter=4 * mm)
    subtitle = ParagraphStyle("CoverSubtitle", parent=styles["Normal"], fontName="MalgunBold", fontSize=13, leading=20, textColor=BLUE, alignment=TA_CENTER)
    tagline = ParagraphStyle("Tagline", parent=styles["Normal"], fontName="Malgun", fontSize=10, leading=16, textColor=MUTED, alignment=TA_CENTER)
    cover_meta = ParagraphStyle("CoverMeta", parent=styles["Normal"], fontName="Malgun", fontSize=10, leading=17, textColor=INK, alignment=TA_CENTER)

    h1 = ParagraphStyle("H1", parent=styles["Heading1"], fontName="MalgunBold", fontSize=17, leading=23, textColor=NAVY, spaceAfter=5 * mm, keepWithNext=True)
    h2 = ParagraphStyle("H2", parent=styles["Heading2"], fontName="MalgunBold", fontSize=11.5, leading=17, textColor=NAVY, spaceBefore=2 * mm, spaceAfter=2.5 * mm, keepWithNext=True)
    body = ParagraphStyle("Body", parent=styles["BodyText"], fontName="Malgun", fontSize=9.5, leading=15.5, textColor=INK, alignment=TA_JUSTIFY, spaceAfter=4.1 * mm, wordWrap="CJK")
    list_item = ParagraphStyle("ListItem", parent=body, leftIndent=5 * mm, spaceAfter=2 * mm, alignment=TA_LEFT, bulletIndent=0)

    story: list = [Spacer(1, 62 * mm)]
    story.append(Paragraph("개발 프로젝트 보고서", kicker_cover))
    story.append(Spacer(1, 6 * mm))
    story.append(Paragraph("MJC AI Campus Agent", title))
    story.append(Paragraph("명지전문대학교 학생을 위한 통합형 AI 캠퍼스 서비스", subtitle))
    story.append(Spacer(1, 4 * mm))
    story.append(Paragraph("학교 정보 접근 격차 해소 · 시간표 및 강의실 안내 · 진로상담 AI 분석", tagline))
    story.append(Spacer(1, 30 * mm))
    story.append(Paragraph("CodeOneBite Team | 2026년 8월", cover_meta))
    story.append(Spacer(1, 5 * mm))
    story.append(Paragraph(
        'GitHub  <link href="https://github.com/hoo419/mjc-ai-hackathon-codeonebite-team" color="#1677FF"><u>github.com/hoo419/mjc-ai-hackathon-codeonebite-team</u></link>',
        cover_meta,
    ))
    story.append(NextPageTemplate("Body"))
    story.append(PageBreak())

    for index, (section_title, blocks) in enumerate(sections):
        story.append(Paragraph(section_title, h1))

        for kind, payload in blocks:
            if kind == "text":
                text = str(payload)
                if text:
                    story.append(Paragraph(text.replace("&", "&amp;"), body))
            elif kind == "h2":
                story.append(Paragraph(str(payload).replace("&", "&amp;"), h2))
            elif kind == "bullets":
                for item in payload:
                    story.append(Paragraph(f"•&nbsp;&nbsp;{item.replace('&', '&amp;')}", list_item))
                story.append(Spacer(1, 2 * mm))
            elif kind == "olist":
                for i, item in enumerate(payload, start=1):
                    story.append(Paragraph(f"{i}.&nbsp;&nbsp;{item.replace('&', '&amp;')}", list_item))
                story.append(Spacer(1, 2 * mm))
            elif kind == "table":
                rows = payload
                if rows:
                    story.append(Spacer(1, 1 * mm))
                    story.append(build_table(rows))
                    story.append(Spacer(1, 4 * mm))

        if index < len(sections) - 1:
            story.append(PageBreak())

    doc.build(story)


def set_run_font(run, name: str, size: float, bold: bool = False, color: str = "172033") -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_docx_hyperlink(paragraph, url: str, text: str, size: float, color: str = "1677FF") -> None:
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:eastAsia"), "맑은 고딕")
    rpr.append(rfonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size * 2)))
    rpr.append(sz)
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), color)
    rpr.append(color_el)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rpr.append(u)
    run.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])


def add_docx_table(doc: Document, rows: list[list[str]]) -> None:
    n_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(cell_text)
            if r_idx == 0:
                set_run_font(run, "맑은 고딕", 9.5, True, "071B33")
            else:
                set_run_font(run, "맑은 고딕", 9.5, False, "172033")
    doc.add_paragraph()


def build_docx(header: list[str], sections: list[tuple[str, list[Block]]]) -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    normal = doc.styles["Normal"]
    normal.font.name = "맑은 고딕"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 17, "071B33", 16, 8),
        ("Heading 2", 12, "1677FF", 12, 6),
    ]:
        style = doc.styles[style_name]
        style.font.name = "맑은 고딕"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(80)
    r = p.add_run("개발 프로젝트 보고서")
    set_run_font(r, "맑은 고딕", 11, True, "1677FF")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("MJC AI Campus Agent")
    set_run_font(r, "맑은 고딕", 28, True, "071B33")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("명지전문대학교 학생을 위한 통합형 AI 캠퍼스 서비스")
    set_run_font(r, "맑은 고딕", 13, True, "1677FF")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("학교 정보 접근 격차 해소 · 시간표 및 강의실 안내 · 진로상담 AI 분석")
    set_run_font(r, "맑은 고딕", 10, False, "5D6B82")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(120)
    r = p.add_run("CodeOneBite Team | 2026년 8월")
    set_run_font(r, "맑은 고딕", 10, False, "172033")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("GitHub  ")
    set_run_font(r, "맑은 고딕", 10, False, "172033")
    add_docx_hyperlink(p, "https://github.com/hoo419/mjc-ai-hackathon-codeonebite-team", "https://github.com/hoo419/mjc-ai-hackathon-codeonebite-team", 10)
    doc.add_page_break()

    for idx, (title, blocks) in enumerate(sections):
        doc.add_heading(title, level=1)
        for kind, payload in blocks:
            if kind == "text":
                text = str(payload)
                if text:
                    p = doc.add_paragraph(text)
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            elif kind == "h2":
                doc.add_heading(str(payload), level=2)
            elif kind == "bullets":
                for item in payload:
                    doc.add_paragraph(item, style="List Bullet")
            elif kind == "olist":
                for item in payload:
                    doc.add_paragraph(item, style="List Number")
            elif kind == "table":
                rows = payload
                if rows:
                    add_docx_table(doc, rows)
        if idx < len(sections) - 1:
            doc.add_page_break()

    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_p.add_run("MJC AI Campus Agent · 개발 프로젝트 보고서  ·  ")
    add_page_number(footer_p)
    for run in footer_p.runs:
        set_run_font(run, "맑은 고딕", 8, False, "5D6B82")
    doc.save(DOCX_OUT)


def main() -> None:
    header, sections = parse_content()
    if len(sections) != 14:
        raise ValueError(f"예상 섹션 수 14개(요약+1~12+부록A), 실제 {len(sections)}개")
    build_pdf(header, sections)
    build_docx(header, sections)
    print(PDF_OUT)
    print(DOCX_OUT)


if __name__ == "__main__":
    main()
